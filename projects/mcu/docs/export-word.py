"""
Export MCU Documentation to Word Format

Converts index.html to a Word document with embedded flowcharts.

Requirements:
    pip install python-docx beautifulsoup4 selenium pillow

Usage:
    python export-word.py

Outputs:
    MCU_Documentation.docx (in same directory)
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
from selenium import webdriver
from selenium.webdriver.common.by import By
from PIL import Image
import time
import io
import os


def create_word_document():
    """Create Word document with Calibri font and numbered headings."""
    doc = Document()

    # Set default font to Calibri
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    return doc


def export_flowcharts_to_images(html_path):
    """
    Render Mermaid flowcharts from HTML and export as PNG images.

    Returns:
        List of image file paths
    """
    print("Launching browser to render flowcharts...")

    # Setup headless Chrome
    options = webdriver.ChromeOptions()
    options.add_argument('--headless')
    options.add_argument('--window-size=1920,1080')
    driver = webdriver.Chrome(options=options)

    # Load HTML
    driver.get(f'file:///{os.path.abspath(html_path)}')
    time.sleep(5)  # Wait for Mermaid to render

    # Find all flowchart containers
    flowcharts = driver.find_elements(By.CSS_SELECTOR, '.flowchart')
    image_paths = []

    for i, flowchart in enumerate(flowcharts):
        # Take screenshot of flowchart
        filename = f'flowchart-{i+1}.png'
        flowchart.screenshot(filename)
        image_paths.append(filename)
        print(f"  Exported {filename}")

    driver.quit()
    print(f"✓ Exported {len(image_paths)} flowcharts\n")

    return image_paths


def parse_html_to_word(html_path, doc):
    """
    Parse HTML and convert to Word document structure.

    Handles:
    - Headers (H2, H3, H4)
    - Paragraphs
    - Tables
    - Lists
    - Code blocks
    - Stats cards
    """
    with open(html_path, 'r', encoding='utf-8') as f:
        html = f.read()

    soup = BeautifulSoup(html, 'html.parser')
    main = soup.find('main')

    if not main:
        raise ValueError("No <main> tag found in HTML")

    # Add title
    title = doc.add_heading('MCU - Margin Call Upload System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Implementation Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0].font
    subtitle_format.size = Pt(14)
    subtitle_format.color.rgb = RGBColor(100, 116, 139)

    doc.add_page_break()

    # Add table of contents placeholder
    toc = doc.add_paragraph('Table of Contents')
    toc.style = 'Heading 1'
    doc.add_paragraph('[Auto-generate in Word: References → Table of Contents → Automatic]')
    doc.add_page_break()

    # Process content
    flowchart_index = 0

    for element in main.find_all(['section', 'h2', 'h3', 'h4', 'p', 'table', 'ul', 'ol', 'pre', 'div']):
        if element.name == 'section':
            continue  # Skip section containers

        elif element.name == 'h2':
            # Heading 1 (numbered)
            doc.add_heading(element.get_text(strip=True), 1)

        elif element.name == 'h3':
            # Heading 2 (numbered)
            doc.add_heading(element.get_text(strip=True), 2)

        elif element.name == 'h4':
            # Heading 3 (numbered)
            doc.add_heading(element.get_text(strip=True), 3)

        elif element.name == 'p':
            # Paragraph
            text = element.get_text(strip=True)
            if text:
                para = doc.add_paragraph(text)
                para.style = 'Normal'

        elif element.name == 'table':
            # Convert HTML table to Word table
            rows = element.find_all('tr')
            if not rows:
                continue

            # Count columns
            cols = len(rows[0].find_all(['th', 'td']))
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = 'Light Grid Accent 1'

            for i, row in enumerate(rows):
                cells = row.find_all(['th', 'td'])
                for j, cell in enumerate(cells):
                    table.rows[i].cells[j].text = cell.get_text(strip=True)

                    # Bold header row
                    if i == 0:
                        for paragraph in table.rows[i].cells[j].paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True

            doc.add_paragraph()  # Spacing after table

        elif element.name in ['ul', 'ol']:
            # Lists
            items = element.find_all('li', recursive=False)
            for item in items:
                text = item.get_text(strip=True)
                para = doc.add_paragraph(text, style='List Bullet' if element.name == 'ul' else 'List Number')

        elif element.name == 'pre':
            # Code block
            code = element.get_text()
            para = doc.add_paragraph(code)
            para.style = 'No Spacing'
            font = para.runs[0].font
            font.name = 'Courier New'
            font.size = Pt(9)

            # Light gray background (simulated with border)
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.right_indent = Inches(0.5)

        elif element.name == 'div' and 'flowchart' in element.get('class', []):
            # Flowchart - insert image
            if flowchart_index < len(flowchart_images):
                doc.add_paragraph(f'[Flowchart {flowchart_index + 1}]')
                doc.add_picture(flowchart_images[flowchart_index], width=Inches(6))
                flowchart_index += 1
                doc.add_paragraph()

        elif element.name == 'div' and 'stats-grid' in element.get('class', []):
            # Stats cards - convert to table
            cards = element.find_all(class_='stat-card')
            if cards:
                table = doc.add_table(rows=1, cols=len(cards))
                table.style = 'Light Grid'

                for i, card in enumerate(cards):
                    value = card.find(class_='stat-value')
                    label = card.find(class_='stat-label')

                    cell_text = f"{value.get_text() if value else ''}\n{label.get_text() if label else ''}"
                    table.rows[0].cells[i].text = cell_text

                doc.add_paragraph()

    print("✓ HTML content converted to Word\n")


def finalize_document(doc, output_path):
    """Add footer and save document."""
    # Add footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "MCU - Margin Call Upload System | Documentation v1.0 | 24 July 2026"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save
    doc.save(output_path)
    print(f"✓ Word document saved: {output_path}")


def cleanup_temp_files(image_paths):
    """Delete temporary flowchart images."""
    for path in image_paths:
        if os.path.exists(path):
            os.remove(path)
    print(f"✓ Cleaned up {len(image_paths)} temporary files")


if __name__ == '__main__':
    HTML_PATH = 'index.html'
    OUTPUT_PATH = 'MCU_Documentation.docx'

    print("=" * 60)
    print("MCU Documentation → Word Export")
    print("=" * 60)
    print()

    if not os.path.exists(HTML_PATH):
        print(f"✗ Error: {HTML_PATH} not found")
        print("  Run this script from projects/mcu/docs/ directory")
        exit(1)

    try:
        # Step 1: Create Word document
        print("1. Creating Word document...")
        doc = create_word_document()

        # Step 2: Export flowcharts as images
        print("2. Rendering flowcharts...")
        flowchart_images = export_flowcharts_to_images(HTML_PATH)

        # Step 3: Parse HTML and convert to Word
        print("3. Converting HTML to Word...")
        parse_html_to_word(HTML_PATH, doc)

        # Step 4: Save document
        print("4. Finalizing document...")
        finalize_document(doc, OUTPUT_PATH)

        # Step 5: Cleanup
        print("5. Cleaning up...")
        cleanup_temp_files(flowchart_images)

        print()
        print("=" * 60)
        print("✓ SUCCESS")
        print("=" * 60)
        print(f"Word document: {os.path.abspath(OUTPUT_PATH)}")
        print()
        print("Next steps:")
        print("1. Open in Word")
        print("2. References → Table of Contents → Automatic Table")
        print("3. Save as final version")

    except Exception as e:
        print(f"\n✗ Error: {e}")
        print("\nTroubleshooting:")
        print("- Install: pip install python-docx beautifulsoup4 selenium pillow")
        print("- Install ChromeDriver: https://chromedriver.chromium.org/")
        exit(1)
